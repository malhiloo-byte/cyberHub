from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/reports/cyberos-complete-project-report-ar.md"
OUTPUT = ROOT / "docs/reports/cyberos-complete-project-report-ar.docx"


def clean_inline(value: str) -> str:
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = value.replace("**", "").replace("__", "")
    return value


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Shading Accent 1"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = clean_inline(value.strip())
    for row in rows[2:] if len(rows) > 1 and all("---" in cell for cell in rows[1]) else rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = clean_inline(value.strip())


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("التقرير الشامل لمشروع CyberOS")
    run.bold = True
    run.font.size = Pt(22)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Personal Cybersecurity Engineering OS\nمن الصفر حتى checkpoint 699fe67b").italic = True

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip() or line.strip() == "---":
            index += 1
            continue
        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=3)
        elif line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=2)
        elif line.startswith("# "):
            document.add_heading(clean_inline(line[2:]), level=1)
        elif line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph()
            paragraph.style = "No Spacing"
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        elif line.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            paragraph.add_run(clean_inline(line[2:]))
        elif line.startswith("| "):
            table_lines: list[list[str]] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append([cell for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(document, table_lines)
            continue
        elif line.startswith("- ") or re.match(r"^\d+\. ", line):
            text = re.sub(r"^(-|\d+\.)\s+", "", line)
            document.add_paragraph(clean_inline(text), style="List Bullet")
        else:
            document.add_paragraph(clean_inline(line))
        index += 1

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
